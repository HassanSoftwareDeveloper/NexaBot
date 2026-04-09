#document_processor.py

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import pandas as pd
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import re
import json
import os
import requests
import time
import logging
from backend.config import settings
from backend.models import ProductInfo
from backend.services.embedding_service import embedding_service

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.products_cache = []
        self.document_types = {}
        
        # Configuration
        self.chunk_size = 800  # Characters per chunk
        self.chunk_overlap = 100
        self.min_chunk_size = 100
        
        # Groq API configuration
        self.groq_api_key = os.getenv("GROQ_API_KEY", "")
        self.groq_url = "https://api.groq.com/openai/v1/chat/completions"
        
        if self.groq_api_key and len(self.groq_api_key) > 20:
            logger.info(f"Groq API configured: {self.groq_api_key[:15]}...")
        else:
            logger.warning("No Groq API key found! Using fallback extraction")
    
    def process_document(self, file_path: Path) -> Dict:
        """
        Process ANY document type and extract information
        Returns: Dict with 'chunks', 'products', 'metadata'
        """
        
        logger.info(f"Processing: {file_path.name}")
        logger.info("=" * 60)
        
        file_ext = file_path.suffix.lower()
        
        try:
            # Extract text based on file type
            if file_ext == '.pdf':
                text = self._extract_pdf_text_enhanced(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_docx_text_complete(file_path)
            elif file_ext in ['.xlsx', '.xls', '.csv']:
                text = self._extract_excel_text(file_path)
            elif file_ext == '.txt':
                text = self._extract_txt_text(file_path)
            elif file_ext == '.json':
                return self._process_json(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                return {'chunks': [], 'products': [], 'metadata': {}}
            
            if not text.strip():
                logger.error("No text extracted from document!")
                return {'chunks': [], 'products': [], 'metadata': {}}
            
            logger.info(f"Extracted {len(text)} characters")
            
            # Clean text
            text = self._clean_text(text)
            
            # Classify document type
            doc_type = self._classify_document_type(file_path.name, text)
            self.document_types[file_path.name] = doc_type
            
            # Create text chunks for vector store
            chunks = self._create_smart_chunks(text, file_path.name, doc_type)
            logger.info(f"Created {len(chunks)} text chunks")
            
            # Extract products (only for product documents)
            products = []
            if doc_type == 'products':
                logger.info("Extracting products...")
                
                if self.groq_api_key and len(self.groq_api_key) > 20:
                    products = self._extract_products_with_groq(text, file_path.name)
                else:
                    products = self._extract_products_fallback(text, file_path.name)
                
                # Filter out non-products
                products = self._filter_real_products(products)
            
            # Create metadata
            metadata = {
                'filename': file_path.name,
                'file_type': file_ext,
                'document_type': doc_type,
                'total_chars': len(text),
                'total_chunks': len(chunks),
                'total_products': len(products)
            }
            
            logger.info(f"Processing complete: {len(chunks)} chunks, {len(products)} products")
            
            return {
                'chunks': chunks,
                'products': products,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            return {'chunks': [], 'products': [], 'metadata': {}}
    
    def _classify_document_type(self, filename: str, content: str) -> str:
        """
        Classify document: products, business_rules, or company_info
        Uses Groq AI if available, otherwise keyword matching
        """
        
        filename_lower = filename.lower()
        content_sample = content[:3000].lower()
        
        # Try Groq classification first
        if self.groq_api_key and len(self.groq_api_key) > 20:
            try:
                classification = self._classify_with_groq(filename, content_sample)
                if classification:
                    logger.info(f"{filename} -> {classification.upper()} (AI classified)")
                    return classification
            except Exception as e:
                logger.warning(f"Groq classification failed: {e}")
        
        # Fallback to keyword matching
        product_keywords = ['product', 'price', 'catalog', 'item', 'stock', 'inventory', 'specification']
        business_keywords = ['rule', 'policy', 'discount', 'offer', 'promotion', 'terms', 'upsell']
        company_keywords = ['about', 'company', 'contact', 'address', 'location', 'team']
        
        product_score = sum(1 for k in product_keywords if k in content_sample or k in filename_lower)
        business_score = sum(1 for k in business_keywords if k in content_sample or k in filename_lower)
        company_score = sum(1 for k in company_keywords if k in content_sample or k in filename_lower)
        
        scores = {'products': product_score, 'business_rules': business_score, 'company_info': company_score}
        doc_type = max(scores, key=scores.get)
        
        logger.info(f"{filename} -> {doc_type.upper()} (keyword classified)")
        return doc_type
    
    def _classify_with_groq(self, filename: str, content: str) -> Optional[str]:
        """Use Groq AI to classify document type"""
        
        prompt = f"""Classify this document into ONE category: "products", "business_rules", or "company_info"

Filename: {filename}
Content: {content[:800]}

Return ONLY one word: products, business_rules, or company_info"""
        
        response = self._call_groq_api(prompt, max_tokens=20, temperature=0.1)
        
        if response:
            response_clean = response.strip().lower()
            if 'product' in response_clean:
                return 'products'
            elif 'business' in response_clean or 'rule' in response_clean:
                return 'business_rules'
            elif 'company' in response_clean or 'info' in response_clean:
                return 'company_info'
        
        return None
    
    def _extract_pdf_text_enhanced(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods"""
        
        full_text = ""
        
        # Method 1: pdfplumber (best for structured data)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n--- Page {page_num} ---\n{text}\n"
            
            if full_text.strip():
                logger.info(f"PDF extracted using pdfplumber")
                return full_text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        # Method 2: PyMuPDF (fallback)
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text:
                    full_text += f"\n--- Page {page_num} ---\n{text}\n"
            doc.close()
            
            if full_text.strip():
                logger.info(f"PDF extracted using PyMuPDF")
                return full_text
        except Exception as e:
            logger.error(f"PyMuPDF failed: {e}")
        
        # Method 3: PyPDF2 (last resort)
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    full_text += page.extract_text() + "\n\n"
            
            if full_text.strip():
                logger.info(f"PDF extracted using PyPDF2")
                return full_text
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
        
        return full_text
    
    def _extract_docx_text_complete(self, file_path: Path) -> str:
        """
        Extract text from BOTH paragraphs AND tables in DOCX
        """
        try:
            doc = Document(file_path)
            all_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text.strip())
            
            # Extract tables (CRITICAL!)
            for table_num, table in enumerate(doc.tables, 1):
                all_text.append(f"\n--- Table {table_num} ---")
                
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        all_text.append(row_text)
            
            full_text = "\n".join(all_text)
            
            logger.info(f"DOCX: {len(doc.paragraphs)} paragraphs + {len(doc.tables)} tables")
            return full_text
            
        except ImportError:
            logger.error("python-docx not installed. Install: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    def _extract_excel_text(self, file_path: Path) -> str:
        """Extract text from Excel/CSV"""
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
                return df.to_string(index=False)
            else:
                full_text = ""
                excel_file = pd.ExcelFile(file_path)
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    full_text += f"\n=== Sheet: {sheet_name} ===\n"
                    full_text += df.to_string(index=False) + "\n"
                return full_text
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            return ""
    
    def _process_json(self, file_path: Path) -> Dict:
        """Process JSON file (likely contains products)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            products = []
            chunks = []
            
            # Handle list of products
            if isinstance(data, list):
                for item in data:
                    if self._is_product_data(item):
                        products.append(item)
                        chunk_text = self._product_to_text(item)
                        chunks.append({
                            'text': chunk_text,
                            'source': file_path.name,
                            'chunk_id': len(chunks),
                            'document_type': 'products'
                        })
            
            # Handle single product or nested structure
            elif isinstance(data, dict):
                if self._is_product_data(data):
                    products.append(data)
                    chunk_text = self._product_to_text(data)
                    chunks.append({
                        'text': chunk_text,
                        'source': file_path.name,
                        'chunk_id': 0,
                        'document_type': 'products'
                    })
                else:
                    # Check nested structure
                    for key, value in data.items():
                        if isinstance(value, list):
                            for item in value:
                                if self._is_product_data(item):
                                    products.append(item)
                                    chunk_text = self._product_to_text(item)
                                    chunks.append({
                                        'text': chunk_text,
                                        'source': file_path.name,
                                        'chunk_id': len(chunks),
                                        'document_type': 'products'
                                    })
            
            metadata = {
                'filename': file_path.name,
                'file_type': '.json',
                'document_type': 'products',
                'total_products': len(products),
                'total_chunks': len(chunks)
            }
            
            logger.info(f"JSON: {len(products)} products extracted")
            
            return {'chunks': chunks, 'products': products, 'metadata': metadata}
            
        except Exception as e:
            logger.error(f"JSON processing failed: {e}")
            return {'chunks': [], 'products': [], 'metadata': {}}
    
    def _is_product_data(self, data: Dict) -> bool:
        """Check if dictionary represents product data"""
        if not isinstance(data, dict) or 'name' not in data:
            return False
        
        product_fields = ['price', 'description', 'category', 'colors', 'sizes', 'shop']
        has_product_fields = sum(1 for field in product_fields if field in data)
        
        return has_product_fields >= 1
    
    def _product_to_text(self, product: Dict) -> str:
        """Convert product dict to searchable text"""
        text_parts = []
        
        if 'name' in product:
            text_parts.append(f"Product: {product['name']}")
        if 'description' in product and product['description']:
            text_parts.append(f"Description: {product['description']}")
        if 'category' in product and product['category']:
            text_parts.append(f"Category: {product['category']}")
        if 'price' in product and product['price']:
            text_parts.append(f"Price: Rs {product['price']}")
        if 'colors' in product and product['colors']:
            text_parts.append(f"Colors: {', '.join(product['colors'][:10])}")
        if 'sizes' in product and product['sizes']:
            text_parts.append(f"Sizes: {', '.join(product['sizes'])}")
        if 'shop' in product and product['shop']:
            text_parts.append(f"Shop: {product['shop']}")
        
        return '. '.join(text_parts) + '.'
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _create_smart_chunks(self, text: str, source: str, doc_type: str) -> List[Dict]:
        """Create intelligent text chunks with overlap"""
        
        if not text or len(text) < self.min_chunk_size:
            if text:
                return [{
                    'text': text,
                    'source': source,
                    'chunk_id': 0,
                    'document_type': doc_type
                }]
            return []
        
        chunks = []
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        current_chunk = ""
        chunk_id = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size:
                if len(current_chunk) >= self.min_chunk_size:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'source': source,
                        'chunk_id': chunk_id,
                        'document_type': doc_type
                    })
                    chunk_id += 1
                    
                    # Add overlap
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk += " " + sentence
            else:
                current_chunk += " " + sentence
        
        # Add final chunk
        if len(current_chunk) >= self.min_chunk_size:
            chunks.append({
                'text': current_chunk.strip(),
                'source': source,
                'chunk_id': chunk_id,
                'document_type': doc_type
            })
        
        return chunks
    
    def _extract_products_with_groq(self, text: str, source: str) -> List[Dict]:
        """Extract products using Groq AI"""
        
        logger.info("Using Groq AI for product extraction...")
        
        all_products = []
        
        # Split text into chunks
        max_chars = 10000
        text_chunks = [text[i:i+max_chars] for i in range(0, len(text), max_chars)]
        
        for chunk_num, text_chunk in enumerate(text_chunks, 1):
            logger.info(f"   Processing chunk {chunk_num}/{len(text_chunks)}...")
            
            products = self._extract_from_chunk_groq(text_chunk, source)
            all_products.extend(products)
            
            # Rate limiting
            if chunk_num < len(text_chunks):
                time.sleep(1)
        
        # Remove duplicates
        unique_products = self._remove_duplicates(all_products)
        
        logger.info(f"Groq extracted {len(unique_products)} unique products")
        return unique_products
    
    def _extract_from_chunk_groq(self, text_chunk: str, source: str) -> List[Dict]:
        """Extract products from a text chunk using Groq"""
        
        prompt = f"""Extract ONLY REAL PRODUCTS from this document.

CRITICAL RULES:
1. Extract ONLY products customers can buy (NOT document headers/titles)
2. SKIP: "Product Name", "Size / Quantity", "Finish(s)", "Business Rules", etc.
3. Extract exact product names, prices, sizes, colors, descriptions
4. Look for prices in formats: Rs 1,200, Rs. 1500, ₹ 2,000, etc.
5. ALL products are in stock (in_stock: true)
6. Return JSON array ONLY

SKIP THESE (NOT products):
- Table headers
- Document titles/sections  
- Business policies
- Rules and procedures
- Company information

Document text:
{text_chunk}

Return JSON array:
[{{"name": "Product Name", "description": "...", "category": "...", "price": 1200.50, "sizes": ["1L", "4L"], "colors": ["Red"], "in_stock": true, "shop": "Bright Paints"}}]

JSON array:"""
        
        response = self._call_groq_api(prompt, max_tokens=3000, temperature=0.1)
        
        if not response:
            return []
        
        try:
            # Clean response
            response_clean = response.strip()
            response_clean = re.sub(r'```json\s*', '', response_clean)
            response_clean = re.sub(r'```\s*$', '', response_clean)
            response_clean = response_clean.strip()
            
            # Parse JSON
            products_data = json.loads(response_clean)
            
            if not isinstance(products_data, list):
                return []
            
            # Convert to dict format
            products = []
            for item in products_data:
                try:
                    # Extract price
                    price = item.get('price')
                    if isinstance(price, str):
                        # Handle price strings like "Rs 1,200", "₹ 2,000", "1500"
                        price = re.sub(r'[^\d.]', '', price)
                    price = float(price) if price else None
                    
                    # Handle in_stock field - ensure it's always a boolean
                    in_stock = item.get('in_stock')
                    if in_stock is None or in_stock == '':
                        in_stock = True  # Default to True if not provided
                    elif isinstance(in_stock, str):
                        in_stock = in_stock.lower() in ['true', 'yes', '1', 'available']
                    
                    product = {
                        'name': str(item.get('name', '')).strip(),
                        'description': str(item.get('description', '')).strip(),
                        'category': str(item.get('category', '')).strip(),
                        'price': price,
                        'colors': item.get('colors', []) if isinstance(item.get('colors'), list) else [],
                        'sizes': item.get('sizes', []) if isinstance(item.get('sizes'), list) else [],
                        'in_stock': bool(in_stock),  # Ensure it's always a boolean
                        'shop': item.get('shop', 'Bright Paints')  # Add shop field with default
                    }
                    
                    if product['name']:
                        products.append(product)
                
                except Exception as e:
                    logger.warning(f"Skipping invalid product: {e}")
                    continue
            
            return products
        
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing failed: {e}")
            return []
        except Exception as e:
            logger.error(f"Product extraction failed: {e}")
            return []
    
    def _call_groq_api(self, prompt: str, max_tokens: int = 2000, temperature: float = 0.1) -> Optional[str]:
        """Call Groq API with retry logic"""
        
        if not self.groq_api_key:
            return None
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.groq_api_key}"
        }
        
        data = {
            "model": "llama-3.3-70b-versatile",
            "messages": [
                {
                    "role": "system",
                    "content": "You extract ONLY real products from documents. NEVER extract document metadata, headers, or business rules. Always return valid JSON with shop field included and in_stock as boolean true."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        
        for attempt in range(3):
            try:
                response = requests.post(self.groq_url, headers=headers, json=data, timeout=30)
                
                if response.status_code == 200:
                    result = response.json()
                    return result['choices'][0]['message']['content']
                elif response.status_code == 429:
                    wait_time = 2 ** attempt
                    logger.warning(f"Rate limited, waiting {wait_time}s...")
                    time.sleep(wait_time)
                    continue
                else:
                    logger.error(f"API error {response.status_code}")
                    return None
            
            except requests.exceptions.Timeout:
                logger.warning(f"Timeout (attempt {attempt + 1}/3)")
                if attempt < 2:
                    time.sleep(2)
                continue
            except Exception as e:
                logger.error(f"API call failed: {e}")
                return None
        
        return None
    
    def _extract_products_fallback(self, text: str, source: str) -> List[Dict]:
        """Fallback product extraction without AI"""
        
        logger.info("Using fallback extraction (no AI)")
        
        products = []
        
        # Pattern: "Product Name - Rs 5,000"
        price_pattern = r'([A-Z][A-Za-z\s]{3,50})\s*[-–—]\s*Rs\.?\s*([\d,]+)'
        matches = re.finditer(price_pattern, text)
        
        for match in matches:
            name = match.group(1).strip()
            price_str = match.group(2).replace(',', '')
            
            try:
                price = float(price_str)
                products.append({
                    'name': name,
                    'price': price,
                    'description': '',
                    'category': '',
                    'colors': [],
                    'sizes': [],
                    'in_stock': True,  # Always True for fallback
                    'shop': 'Bright Paints'  # Add shop field
                })
            except ValueError:
                continue
        
        return products
    
    def _filter_real_products(self, products: List[Dict]) -> List[Dict]:
        """Filter out non-products (metadata, headers, titles)"""
        
        non_product_keywords = [
            'business rules', 'executive summary', 'company overview',
            'sales & distribution', 'pricing & discounts', 'ordering',
            'inventory', 'manufacturing', 'returns', 'warranty',
            'health, safety', 'protective equipment', 'price on request',
            'product name', 'size / quantity', 'finish(s)', 'purpose',
            'description', 'quality', 'special', 'promotional', 'fifo',
            'cycle counts', 'ppe mandatory'
        ]
        
        real_products = []
        
        for product in products:
            name = product.get('name', '').lower()
            
            # Skip if contains non-product keywords
            if any(keyword in name for keyword in non_product_keywords):
                logger.info(f"   Skipping: {product.get('name', '')}")
                continue
            
            # Skip if too long (likely a title)
            if len(product.get('name', '')) > 100:
                continue
            
            # Skip if contains em dash
            if '—' in product.get('name', '') or '–' in product.get('name', ''):
                continue
            
            # Skip if starts with number
            if re.match(r'^\d+\.', product.get('name', '')):
                continue
            
            real_products.append(product)
        
        logger.info(f"Filtered to {len(real_products)} real products")
        return real_products
    
    def _remove_duplicates(self, products: List[Dict]) -> List[Dict]:
        """Remove duplicate products"""
        seen_names = set()
        unique = []
        
        for product in products:
            name = product.get('name', '').lower().strip()
            if name and name not in seen_names:
                seen_names.add(name)
                unique.append(product)
        
        return unique
    
    def process_multiple_documents(self, file_paths: List[Path]) -> Dict:
        """Process multiple documents"""
        all_chunks = []
        all_products = []
        metadata_list = []
        
        for file_path in file_paths:
            result = self.process_document(file_path)
            all_chunks.extend(result['chunks'])
            all_products.extend(result['products'])
            metadata_list.append(result['metadata'])
        
        logger.info(f"Total: {len(file_paths)} files, {len(all_chunks)} chunks, {len(all_products)} products")
        
        return {
            'chunks': all_chunks,
            'products': all_products,
            'metadata': {
                'total_files': len(file_paths),
                'total_chunks': len(all_chunks),
                'total_products': len(all_products),
                'files': metadata_list
            }
        }
    
    def save_products_to_json(self, products: List[Dict], filename: str):
        """Save products to JSON"""
        if not products:
            logger.warning("No products to save")
            return None
        
        settings.products_dir.mkdir(parents=True, exist_ok=True)
        output_path = settings.products_dir / f"{filename}.json"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(products, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved {len(products)} products to {output_path.name}")
        return output_path

# Global instance
document_processor = DocumentProcessor()